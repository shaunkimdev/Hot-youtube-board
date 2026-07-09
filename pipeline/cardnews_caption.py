from docx import Document
from docx.oxml.ns import qn

FONT = "맑은 고딕"


def set_korean_font(run, name=FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = rPr.makeelement(qn('w:lang'), {})
        rPr.append(lang)
    lang.set(qn('w:eastAsia'), 'ko-KR')


def add_para(doc, text=""):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_korean_font(run)
    return para


def save_caption(path, lines):
    """lines: list of strings; each string is one paragraph (use '' for blank line)."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT
    rPr = style.element.get_or_add_rPr() if hasattr(style.element, 'get_or_add_rPr') else None
    for text in lines:
        add_para(doc, text)
    doc.save(path)
    print("saved", path)
